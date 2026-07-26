from __future__ import annotations

import csv
import argparse
import json
import shutil
import zipfile
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from screen_layout_contracts import build_layout_catalog


ROOT = Path(__file__).resolve().parent
SNAPSHOT = ROOT / "full_snapshot.json"
OUT = Path.home() / "Downloads" / "Resonance_Carbonet_전수설계문서_20260723_v2.0"
ZIP = OUT.parent / f"{OUT.name}.zip"
PROJECT = "Resonance / Carbonet 운영 시스템 전수 설계 기준서"
TODAY = "2026-07-23"
FONT = "Malgun Gothic"
CONTENT_DXA = 9360
NAVY = "0B2545"
BLUE = "2E74B5"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
GREEN = "216E39"
GOLD = "7A5A00"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def text(value) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, separators=(",", ":"))
    return str(value)


def parse_json_value(value):
    if value in (None, ""):
        return []
    if isinstance(value, (list, dict)):
        return value
    try:
        return json.loads(value)
    except Exception:
        return value


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=80, bottom=55, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError((widths, sum(widths)))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
    if tbl_w.getparent() is None:
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
    if tbl_ind.getparent() is None:
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "80")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
            if tc_w.getparent() is None:
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure_doc(doc: Document, short_title: str) -> None:
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.72)
    sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.34)
    sec.footer_distance = Inches(0.34)
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb("20252B")
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.12
    for name, size, color, before, after in [
        ("Title", 24, NAVY, 0, 6),
        ("Subtitle", 11, MID_GRAY, 0, 12),
        ("Heading 1", 15, BLUE, 14, 7),
        ("Heading 2", 11.5, NAVY, 10, 5),
        ("Heading 3", 10, NAVY, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = sec.header.paragraphs[0]
    header.add_run(f"{PROJECT}  |  {short_title}")
    for run in header.runs:
        run.font.name = FONT
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        run.font.size = Pt(7.5)
        run.font.color.rgb = rgb(MID_GRAY)
    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("LIVE SNAPSHOT FULL INVENTORY  |  ")
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(7.5)
    run.font.color.rgb = rgb(MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)


def add_table(doc: Document, headers: list[str], rows, widths: list[int], font_size=7.3) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for idx, value in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = value
        set_cell_shading(cell, LIGHT_BLUE)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = rgb(NAVY)
            run.font.size = Pt(font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = text(value)
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.name = FONT
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
                    run.font.size = Pt(font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_callout(doc: Document, label: str, value: str, fill=LIGHT_BLUE, color=NAVY) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = rgb(color)
    p.add_run(value)
    set_table_geometry(table, [CONTENT_DXA])


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.32)
        p.paragraph_format.first_line_indent = Inches(-0.16)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(item)


def new_doc(title: str, doc_id: str, purpose: str, snap: dict) -> Document:
    doc = Document()
    configure_doc(doc, title)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("SYSTEM DESIGN DELIVERABLE · FULL INVENTORY")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = rgb(BLUE)
    doc.add_paragraph(title, style="Title")
    doc.add_paragraph(PROJECT, style="Subtitle")
    add_table(doc, ["항목", "값"], [
        ["문서 ID", doc_id],
        ["버전/기준일", f"v2.0 / {TODAY}"],
        ["스냅샷", snap["captured_at"]],
        ["대상", f"{snap['host']} / {snap['namespace']} / DB {snap['database']}"],
        ["목적", purpose],
    ], [1900, 7460], 8.3)
    add_callout(
        doc,
        "포함 기준",
        "운영 DB 및 서버 소스에서 자동 추출 가능한 현재 자산을 전수 수록했습니다. 사람의 업무 의미가 원천 자산에 없거나 검증 플래그가 false인 항목은 추정하지 않고 검토 필요 상태로 유지합니다.",
        "FFF5D6",
        GOLD,
    )
    doc.add_page_break()
    return doc


def add_overview(doc: Document, snap: dict, specific: list[list[str]]) -> None:
    doc.add_heading("1. 전수 범위와 기준", level=1)
    db = snap["database_summary"]
    sc = snap["screen_counts"]
    rows = [
        ["화면 계약", sc["contract_rows"], f"고유 route {sc['contract_routes']}개"],
        ["DB", db["table_count"], f"컬럼 {db['column_count']} / 인덱스 {db['index_count']} / 제약 {db['constraint_count']}"],
        ["API 소스 매핑", len(snap["api_mappings"]), "Spring mapping 어노테이션 라인"],
        ["테스트 자산", len(snap["test_files"]), "테스트/스펙/E2E 파일"],
        ["마이그레이션", len(snap["migration_files"]), "Flyway/Liquibase/SQL·XML"],
        *specific,
    ]
    add_table(doc, ["자산", "수량", "비고"], rows, [2200, 1300, 5860], 8.0)
    add_bullets(doc, [
        "전수는 문서 종류와 직접 관련된 원천 자산의 모든 행/파일/매핑을 뜻합니다.",
        "동일 route의 복수 계약은 삭제하지 않고 계약 ID별로 모두 유지합니다.",
        "DB 품질 테이블이 현재 존재하지 않으므로 품질 점수는 계약 검증 플래그와 별개로 미확인 처리합니다.",
    ])


def save(doc: Document, filename: str) -> Path:
    path = OUT / filename
    doc.save(path)
    return path


def decoded_contracts(snap: dict) -> list[dict]:
    return [json.loads(row["row_json"]) for row in snap["screen_contracts"]]


def decoded_asset_rows(snap: dict, names: list[str]) -> list[list[str]]:
    rows = []
    for name in names:
        for item in snap.get("asset_rows", {}).get(name, []):
            rows.append([name, item["row_json"]])
    return rows


def unique_routes(contracts: list[dict]) -> list[dict]:
    chosen = {}
    for c in contracts:
        chosen[c.get("route_path")] = c
    return [chosen[k] for k in sorted(chosen, key=lambda x: x or "")]


def pattern(route: str) -> str:
    if route in ("/", "/home") or route.endswith("/"):
        return "대시보드/홈"
    tail = route.rstrip("/").split("/")[-1].lower()
    if any(k in tail for k in ("detail", "view", "info")):
        return "상세"
    if any(k in tail for k in ("write", "create", "new", "register", "edit", "modify", "step")):
        return "입력/단계"
    if any(k in tail for k in ("list", "management", "manage", "status", "history")):
        return "목록/관리"
    if any(k in route.lower() for k in ("signin", "join", "findpassword", "findid")):
        return "인증/가입"
    return "업무 기본"


def build_docs(snap: dict) -> list[Path]:
    c = decoded_contracts(snap)
    routes = unique_routes(c)
    layout_catalog = build_layout_catalog(c)
    layouts_by_contract = {x["identity"]["contractId"]: x for x in layout_catalog["contracts"]}
    paths: list[Path] = []

    actors = defaultdict(lambda: {"processes": set(), "routes": set(), "audiences": set()})
    for x in c:
        a = x.get("actor_code") or "UNSPECIFIED"
        actors[a]["processes"].add(x.get("process_code") or "")
        actors[a]["routes"].add(x.get("route_path") or "")
        actors[a]["audiences"].add(x.get("audience") or "")
    d = new_doc("액터 정의서", "ACTOR-FULL-001", "모든 화면 계약 액터와 DB 역할 자산의 책임·범위를 정의합니다.", snap)
    add_overview(d, snap, [["계약 액터", len(actors), "actor_code 전수"]])
    d.add_heading("2. 액터 전수 목록", level=1)
    add_table(d, ["액터", "Audience", "프로세스 수", "route 수", "프로세스"], [
        [a, ", ".join(sorted(v["audiences"])), len(v["processes"]), len(v["routes"]), ", ".join(sorted(v["processes"]))]
        for a, v in sorted(actors.items())
    ], [1500, 1100, 900, 800, 5060], 7.2)
    d.add_heading("3. 역할·권한 원천 행", level=1)
    add_table(d, ["원천 테이블", "전체 행 JSON"], decoded_asset_rows(snap, ["comtnroleinfo", "msatnroleinfo", "msatnroles_hierarchy", "comtnauthorrolerelate", "msatnauthorrolerelate"]), [2400, 6960], 6.8)
    paths.append(save(d, "01_액터_정의서.docx"))

    usecases = defaultdict(lambda: {"actors": set(), "routes": set(), "purposes": set(), "entry": set(), "exit": set()})
    for x in c:
        key = (x.get("process_code") or "UNSPECIFIED", x.get("step_code") or "UNSPECIFIED")
        usecases[key]["actors"].add(x.get("actor_code") or "")
        usecases[key]["routes"].add(x.get("route_path") or "")
        usecases[key]["purposes"].add(x.get("business_purpose") or "")
        usecases[key]["entry"].add(x.get("entry_condition") or "")
        usecases[key]["exit"].add(x.get("exit_condition") or "")
    d = new_doc("유즈케이스 명세서", "UC-FULL-001", "process-step 조합별 모든 액터·화면·선후조건을 정의합니다.", snap)
    add_overview(d, snap, [["유즈케이스", len(usecases), "process_code + step_code"]])
    d.add_heading("2. 유즈케이스 전수 목록", level=1)
    add_table(d, ["Process", "Step", "Actor", "route 수", "업무 목적", "진입", "종료"], [
        [k[0], k[1], ",".join(sorted(v["actors"])), len(v["routes"]), " | ".join(sorted(v["purposes"])), " | ".join(sorted(v["entry"])), " | ".join(sorted(v["exit"]))]
        for k, v in sorted(usecases.items())
    ], [1300, 1200, 1100, 600, 1900, 1600, 1660], 6.7)
    paths.append(save(d, "02_유즈케이스_명세서.docx"))

    d = new_doc("BPMN 설계서", "BPMN-FULL-001", "프로세스 정의·단계·흐름 edge의 전체 원천을 수록합니다.", snap)
    bpmn_rows = decoded_asset_rows(snap, ["framework_process_definition", "framework_process_step", "framework_process_flow_edge"])
    add_overview(d, snap, [["BPMN 원천 행", len(bpmn_rows), "정의/단계/edge"]])
    d.add_heading("2. 프로세스·단계·흐름 전수", level=1)
    add_table(d, ["원천 테이블", "전체 행 JSON"], bpmn_rows, [2600, 6760], 6.6)
    paths.append(save(d, "03_BPMN_설계서.docx"))

    d = new_doc("비즈니스 로직 설계서", "BIZ-LOGIC-FULL-001", "프로세스 순서·데이터 인계·준비도 원천과 화면 진출입 규칙을 전수 수록합니다.", snap)
    logic_rows = decoded_asset_rows(snap, ["framework_business_process_sequence", "framework_process_data_handoff", "framework_process_professional_readiness"])
    add_overview(d, snap, [["로직 원천 행", len(logic_rows), "sequence/handoff/readiness"]])
    d.add_heading("2. 비즈니스 로직 원천", level=1)
    add_table(d, ["원천 테이블", "전체 행 JSON"], logic_rows, [2700, 6660], 6.6)
    d.add_heading("3. 계약 진출입 로직", level=1)
    add_table(d, ["계약", "Process/Step", "Route", "진입 조건", "종료 조건", "상태"], [
        [x.get("contract_id"), f"{x.get('process_code')}/{x.get('step_code')}", x.get("route_path"), x.get("entry_condition"), x.get("exit_condition"), x.get("contract_status")]
        for x in c
    ], [700, 1600, 1900, 1900, 1900, 1360], 6.8)
    paths.append(save(d, "04_비즈니스_로직_설계서.docx"))

    d = new_doc("시나리오 명세서", "SCENARIO-FULL-001", "DB 전문 시나리오와 화면 계약별 업무 시나리오를 모두 수록합니다.", snap)
    scenario_rows = decoded_asset_rows(snap, ["framework_process_professional_scenario"])
    add_overview(d, snap, [["전문 시나리오 행", len(scenario_rows), "DB 원천"], ["계약 시나리오", len(c), "계약 ID별"]])
    d.add_heading("2. 전문 시나리오 원천", level=1)
    add_table(d, ["원천", "전체 행 JSON"], scenario_rows, [2200, 7160], 6.6)
    d.add_heading("3. 화면 계약 시나리오", level=1)
    add_table(d, ["계약", "Actor", "Process/Step", "Route", "목적", "진입", "종료"], [
        [x.get("contract_id"), x.get("actor_code"), f"{x.get('process_code')}/{x.get('step_code')}", x.get("route_path"), x.get("business_purpose"), x.get("entry_condition"), x.get("exit_condition")]
        for x in c
    ], [650, 1000, 1400, 1700, 1900, 1350, 1360], 6.5)
    paths.append(save(d, "05_시나리오_명세서.docx"))

    d = new_doc("요구사항 정의서", "REQ-FULL-001", "모든 화면 계약을 추적 가능한 기능·비기능 요구사항으로 전개합니다.", snap)
    add_overview(d, snap, [["계약 기반 요구사항", len(c), "REQ-CONTRACT-{id}"]])
    d.add_heading("2. 요구사항 전수", level=1)
    add_table(d, ["요구사항 ID", "Actor", "Process/Step", "Route", "요구 내용", "완료 조건", "상태"], [
        [f"REQ-CONTRACT-{x.get('contract_id')}", x.get("actor_code"), f"{x.get('process_code')}/{x.get('step_code')}", x.get("route_path"), x.get("business_purpose"), x.get("exit_condition"), x.get("contract_status")]
        for x in c
    ], [1200, 950, 1400, 1750, 1900, 1250, 910], 6.5)
    paths.append(save(d, "06_요구사항_정의서.docx"))

    d = new_doc("비즈니스 규칙 정의서", "RULE-FULL-001", "계약별 진입·종료·보안·접근성·응답형 규칙을 전수 수록합니다.", snap)
    add_overview(d, snap, [["계약 규칙", len(c), "계약 ID별"]])
    d.add_heading("2. 비즈니스 규칙 전수", level=1)
    add_table(d, ["계약", "Route", "진입/종료", "보안", "접근성", "반응형", "검증"], [
        [x.get("contract_id"), x.get("route_path"), f"IN:{x.get('entry_condition')} / OUT:{x.get('exit_condition')}", x.get("security_contract"), x.get("accessibility_contract"), x.get("responsive_contract"), f"AUTH={x.get('authority_verified')};A11Y={x.get('accessibility_verified')};RESP={x.get('responsive_verified')}"]
        for x in c
    ], [650, 1700, 1900, 1550, 1450, 1250, 860], 6.3)
    paths.append(save(d, "07_비즈니스_규칙_정의서.docx"))

    domains = defaultdict(list)
    for t in snap["tables"]:
        name = t["table_name"]
        prefix = name.split("_")[0] if "_" in name else name[:6]
        domains[prefix].append(name)
    d = new_doc("개념 데이터 모델", "CDM-FULL-001", "전체 테이블을 업무 개념군으로 분류하고 개념 간 관계 원천을 제공합니다.", snap)
    add_overview(d, snap, [["개념군", len(domains), "테이블 명명 규칙 기반"], ["테이블", len(snap["tables"]), "전수"]])
    d.add_heading("2. 개념군 및 소속 엔터티", level=1)
    add_table(d, ["개념군", "엔터티 수", "엔터티 전체"], [[k, len(v), ", ".join(sorted(v))] for k, v in sorted(domains.items())], [1500, 900, 6960], 6.8)
    paths.append(save(d, "08_개념_데이터_모델.docx"))

    fk = [x for x in snap["constraints"] if x["constraint_type"] == "FOREIGN KEY"]
    d = new_doc("논리 데이터 모델", "LDM-FULL-001", "전체 논리 엔터티와 PK/FK/UK 관계를 정의합니다.", snap)
    add_overview(d, snap, [["논리 엔터티", len(snap["tables"]), "테이블/뷰"], ["FK 관계", len(fk), "전수"]])
    d.add_heading("2. 논리 엔터티 전수", level=1)
    add_table(d, ["Schema", "Entity", "유형", "추정 행", "설명"], [[x["table_schema"], x["table_name"], x["table_type"], x["estimated_rows"], x["description"]] for x in snap["tables"]], [1000, 3000, 1100, 900, 3360], 7.0)
    d.add_heading("3. 외래키 관계 전수", level=1)
    add_table(d, ["Schema", "Entity", "FK", "컬럼", "참조 엔터티", "참조 컬럼"], [[x["constraint_schema"], x["table_name"], x["constraint_name"], x["columns"], x["referenced_table"], x["referenced_column"]] for x in fk], [850, 1900, 1900, 1300, 2100, 1310], 6.7)
    paths.append(save(d, "09_논리_데이터_모델.docx"))

    d = new_doc("물리 데이터 모델", "PDM-FULL-001", "전체 물리 테이블·제약·인덱스 구조를 정의합니다.", snap)
    add_overview(d, snap, [["제약조건", len(snap["constraints"]), "전수"], ["인덱스", len(snap["indexes"]), "전수"]])
    d.add_heading("2. 제약조건 전수", level=1)
    add_table(d, ["Schema", "Table", "Type", "Name", "Columns", "Reference"], [[x["constraint_schema"], x["table_name"], x["constraint_type"], x["constraint_name"], x["columns"], f"{x['referenced_table']}.{x['referenced_column']}"] for x in snap["constraints"]], [800, 1700, 1100, 2100, 1500, 2160], 6.4)
    d.add_heading("3. 인덱스 전수", level=1)
    add_table(d, ["Schema", "Table", "Index", "Definition"], [[x["table_schema"], x["table_name"], x["index_name"], x["indexdef"]] for x in snap["indexes"]], [800, 1800, 2100, 4660], 6.4)
    paths.append(save(d, "10_물리_데이터_모델.docx"))

    d = new_doc("테이블 정의서", "TABLE-FULL-001", "344개 테이블의 4,243개 컬럼을 기본값·NULL·설명과 함께 전수 정의합니다.", snap)
    add_overview(d, snap, [["컬럼", len(snap["columns"]), "전수"]])
    d.add_heading("2. 컬럼 정의 전수", level=1)
    add_table(d, ["Schema", "Table", "#", "Column", "Type/Length", "Null", "Default", "Description"], [[x["table_schema"], x["table_name"], x["ordinal_position"], x["column_name"], f"{x['data_type']}({x['max_length']})" if x["max_length"] else x["data_type"], x["is_nullable"], x["column_default"], x["description"]] for x in snap["columns"]], [650, 1550, 400, 1450, 1200, 500, 1750, 1860], 6.2)
    paths.append(save(d, "11_테이블_정의서.docx"))

    d = new_doc("화면 정의서", "SCREEN-FULL-001", "1,181개 계약과 1,000개 고유 route의 화면 메타데이터를 전수 정의합니다.", snap)
    add_overview(d, snap, [["화면 계약", len(c), "전수"], ["고유 route", len(routes), "전수"]])
    d.add_heading("2. 화면 계약 전수", level=1)
    add_table(d, ["계약", "Route", "화면명", "Audience", "Actor", "Process/Step", "목적", "상태"], [[x.get("contract_id"), x.get("route_path"), x.get("screen_name"), x.get("audience"), x.get("actor_code"), f"{x.get('process_code')}/{x.get('step_code')}", x.get("business_purpose"), x.get("contract_status")] for x in c], [600, 1700, 1200, 750, 950, 1350, 1900, 910], 6.3)
    paths.append(save(d, "12_화면_정의서.docx"))

    d = new_doc("와이어프레임", "WIREFRAME-FULL-001", "모든 고유 route의 화면 패턴과 섹션·필드·명령 구조를 텍스트 와이어프레임으로 정의합니다.", snap)
    add_overview(d, snap, [["route 와이어프레임", len(routes), "고유 route 전수"]])
    d.add_heading("2. route별 와이어프레임 구조", level=1)
    add_table(d, ["Route", "화면명", "패턴", "Sections", "Fields", "Commands", "States"], [[x.get("route_path"), x.get("screen_name"), pattern(x.get("route_path") or ""), x.get("section_contract"), x.get("field_contract"), x.get("command_contract"), x.get("state_contract")] for x in routes], [1800, 1300, 900, 1450, 1450, 1450, 1010], 6.2)
    paths.append(save(d, "13_와이어프레임.docx"))

    d = new_doc("UI 상세 설계서", "UI-FULL-001", "계약별 UI 섹션·필드·명령·상태·데이터·접근성·반응형 상세를 전수 정의합니다.", snap)
    add_overview(d, snap, [["UI 계약", len(c), "계약 ID 전수"]])
    d.add_heading("2. UI 상세 계약 전수", level=1)
    add_table(d, ["계약/Route", "Section", "Field", "Command", "State", "Data", "A11y/Responsive"], [[f"{x.get('contract_id')}\n{x.get('route_path')}", x.get("section_contract"), x.get("field_contract"), x.get("command_contract"), x.get("state_contract"), x.get("data_contract"), f"{x.get('accessibility_contract')} / {x.get('responsive_contract')}"] for x in c], [1450, 1250, 1450, 1350, 1050, 1300, 1510], 6.0)
    paths.append(save(d, "14_UI_상세_설계서.docx"))

    d = new_doc("화면 흐름도", "FLOW-FULL-001", "Actor-Process-Step-Route 연결을 계약 단위로 전수 정의합니다.", snap)
    add_overview(d, snap, [["흐름 연결", len(c), "계약 ID 전수"]])
    d.add_heading("2. 화면 흐름 연결 전수", level=1)
    add_table(d, ["계약", "Actor", "Process", "Step", "Route", "Entry", "Exit"], [[x.get("contract_id"), x.get("actor_code"), x.get("process_code"), x.get("step_code"), x.get("route_path"), x.get("entry_condition"), x.get("exit_condition")] for x in c], [600, 1050, 1350, 1200, 1900, 1630, 1630], 6.5)
    paths.append(save(d, "15_화면_흐름도.docx"))

    contract_apis = []
    for x in c:
        api = parse_json_value(x.get("api_contract"))
        if isinstance(api, list):
            for item in api:
                contract_apis.append([x.get("contract_id"), x.get("route_path"), item.get("method", "") if isinstance(item, dict) else "", item.get("path", "") if isinstance(item, dict) else text(item), item.get("desc", "") if isinstance(item, dict) else "", x.get("api_verified")])
        elif api:
            contract_apis.append([x.get("contract_id"), x.get("route_path"), "", text(api), "", x.get("api_verified")])
    d = new_doc("API 설계서", "API-FULL-001", "화면 계약 API와 백엔드 Spring 매핑 원천을 전수 수록합니다.", snap)
    add_overview(d, snap, [["계약 API", len(contract_apis), "api_contract 전개"], ["소스 매핑", len(snap["api_mappings"]), "어노테이션 라인 전수"]])
    d.add_heading("2. 화면 계약 API", level=1)
    add_table(d, ["계약", "화면 Route", "Method", "API Path", "설명", "검증"], contract_apis, [650, 1900, 700, 2200, 3110, 800], 6.6)
    d.add_heading("3. 백엔드 매핑 소스 원천", level=1)
    add_table(d, ["No", "소스 매핑 라인"], [[i + 1, line] for i, line in enumerate(snap["api_mappings"])], [650, 8710], 6.4)
    paths.append(save(d, "16_API_설계서.docx"))

    d = new_doc("테스트 계획서", "TEST-PLAN-FULL-001", "전체 화면·API·DB·프로세스 자산에 대한 계층별 테스트 범위와 실제 테스트 파일을 정의합니다.", snap)
    add_overview(d, snap, [["테스트 파일", len(snap["test_files"]), "전수"], ["검증 필요 계약", sum(1 for x in c if not all([x.get('api_verified'), x.get('database_verified'), x.get('authority_verified'), x.get('responsive_verified'), x.get('accessibility_verified'), x.get('exception_states_verified')])), "계약 검증 플래그 기준"]])
    d.add_heading("2. 테스트 자산 파일 전수", level=1)
    add_table(d, ["No", "테스트 자산 경로"], [[i + 1, x] for i, x in enumerate(snap["test_files"])], [700, 8660], 6.5)
    d.add_heading("3. 품질 게이트", level=1)
    add_bullets(d, ["변경 route smoke 후 전체 route gate를 수행합니다.", "API·DB·권한·반응형·접근성·예외 상태의 여섯 검증 플래그를 계약별 종료 기준으로 사용합니다.", "실제 품질 점수 원천 테이블 부재 때문에 점수 기반 통과를 선언하지 않습니다."])
    paths.append(save(d, "17_테스트_계획서.docx"))

    d = new_doc("테스트 시나리오", "TEST-SCENARIO-FULL-001", "모든 고유 route에 대한 접근·상태·API·권한·접근성 시나리오를 생성합니다.", snap)
    add_overview(d, snap, [["route 시나리오", len(routes), "고유 route 전수"]])
    d.add_heading("2. route별 전수 테스트 시나리오", level=1)
    add_table(d, ["Test ID", "Route", "Actor", "사전조건", "절차", "기대결과", "현재 검증"], [[f"TST-ROUTE-{i:04d}", x.get("route_path"), x.get("actor_code"), x.get("entry_condition"), "인증 세션으로 route 접근 → loading/empty/error/ready 상태 → 계약 API 및 명령 실행 → 접근성 이름·키보드·반응형 확인", x.get("exit_condition"), f"API={x.get('api_verified')};DB={x.get('database_verified')};AUTH={x.get('authority_verified')};RESP={x.get('responsive_verified')};A11Y={x.get('accessibility_verified')};EX={x.get('exception_states_verified')}" ] for i, x in enumerate(routes, 1)], [1050, 1650, 950, 1450, 2100, 1200, 960], 6.2)
    paths.append(save(d, "18_테스트_시나리오.docx"))
    d = new_doc(
        "화면 상세 설계서",
        "SCREEN-DETAIL-FULL-001",
        "고유 route별 기준 좌표, CSS Grid/Flex 배치, 섹션·컴포넌트 영역, 반응형 전환과 상태 표현을 전수 정의합니다.",
        snap,
    )
    add_overview(d, snap, [
        ["화면 상세 설계", len(routes), "고유 route 전수"],
        ["레이아웃 계약", len(layout_catalog["contracts"]), "계약 ID 전수"],
        ["기준 뷰포트", 4, "1440 / 1280 / 768 / 360"],
    ])
    d.add_heading("2. 구현 기준", level=1)
    add_callout(
        d,
        "좌표 사용 원칙",
        "x/y/width/height는 기준 뷰포트에서 검토·시각 회귀를 위한 참조값입니다. 실제 구현의 권위 기준은 DOM 순서와 CSS Grid/Flex 제약이며 일반 화면 흐름에서 absolute positioning을 사용하지 않습니다.",
        "FFF5D6",
        GOLD,
    )
    add_bullets(d, [
        "Desktop은 1440x900, 12열 Grid, 240px sidebar, 64px topbar, 24px gap을 기준으로 합니다.",
        "Tablet 768px부터 sidebar를 접고 8열 Grid, Mobile 360px에서는 4열 단일 흐름으로 전환합니다.",
        "계약에 section/field/command가 있으면 그대로 사용하고, 없으면 route 패턴 템플릿을 사용하되 requiresHumanReview=true로 표시합니다.",
        "LOADING, EMPTY, ERROR, FORBIDDEN, READY 상태는 콘텐츠 영역을 보존하는 overlay 또는 inline 상태로 구현합니다.",
    ])
    d.add_heading("3. route별 화면 상세 배치", level=1)
    detail_rows = []
    for route_contract in routes:
        layout = layouts_by_contract[route_contract.get("contract_id")]
        desktop = layout["responsive"]["desktop"]
        section_geometry = " | ".join(
            f"{x['sectionId']}({x['rect']['x']},{x['rect']['y']},{x['rect']['width']},{x['rect']['height']})"
            for x in desktop["sections"]
        )
        component_geometry = " | ".join(
            f"{x['id']}:{x['component']}@{x['sectionId']}({x['referenceRect']['x']},{x['referenceRect']['y']},{x['referenceRect']['width']},{x['referenceRect']['height']})"
            for x in layout["components"]
        )
        shell = desktop["shell"]["content"]
        responsive = " / ".join(
            f"{name}:{frame['grid']['columns']}col,{frame['shell']['content']['width']}px"
            for name, frame in layout["responsive"].items()
        )
        detail_rows.append([
            f"{layout['identity']['contractId']}\n{layout['identity']['route']}",
            f"{layout['identity']['screenName']}\n{layout['pattern']}",
            f"content({shell['x']},{shell['y']},{shell['width']},{shell['minHeight']})\n12col/gap24",
            section_geometry,
            component_geometry,
            responsive,
            f"{','.join(x['name'] for x in layout['states'])}\nreview={layout['traceability']['requiresHumanReview']}",
        ])
    add_table(
        d,
        ["계약/Route", "화면/패턴", "Shell/Grid", "Section 좌표(x,y,w,h)", "Component 좌표(x,y,w,h)", "Responsive", "States/Review"],
        detail_rows,
        [1350, 1150, 1150, 1650, 1750, 1350, 960],
        5.8,
    )
    d.add_heading("4. 기계 판독 계약", level=1)
    add_callout(
        d,
        "개발 입력",
        "evidence/screen-layout-contracts.json은 동일 내용을 손실 없이 제공하며 development/packets/*.json의 design.layout에도 포함됩니다.",
        "E9F6EC",
        GREEN,
    )
    paths.append(save(d, "19_화면_상세_설계서.docx"))
    return paths


def export_evidence(snap: dict, contracts: list[dict]) -> None:
    evidence = OUT / "evidence"
    evidence.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SNAPSHOT, evidence / "full_snapshot.json")
    layout_catalog = build_layout_catalog(contracts)
    (evidence / "screen-layout-contracts.json").write_text(
        json.dumps(layout_catalog, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    layout_rows = []
    section_rows = []
    component_rows = []
    for layout in layout_catalog["contracts"]:
        ident = layout["identity"]
        desktop = layout["responsive"]["desktop"]
        layout_rows.append({
            "contract_id": ident["contractId"],
            "route": ident["route"],
            "screen_name": ident["screenName"],
            "pattern": layout["pattern"],
            "section_count": len(layout["sections"]),
            "component_count": len(layout["components"]),
            "requires_human_review": layout["traceability"]["requiresHumanReview"],
            "desktop_content_x": desktop["shell"]["content"]["x"],
            "desktop_content_y": desktop["shell"]["content"]["y"],
            "desktop_content_width": desktop["shell"]["content"]["width"],
            "desktop_document_height": desktop["documentHeight"],
        })
        for section in desktop["sections"]:
            section_rows.append({
                "contract_id": ident["contractId"],
                "route": ident["route"],
                "section_id": section["sectionId"],
                **section["rect"],
                "grid_column_start": section["grid"]["columnStart"],
                "grid_column_span": section["grid"]["columnSpan"],
                "layout": section["layout"],
            })
        for component in layout["components"]:
            component_rows.append({
                "contract_id": ident["contractId"],
                "route": ident["route"],
                "component_id": component["id"],
                "label": component["label"],
                "component_type": component["component"],
                "section_id": component["sectionId"],
                **component["referenceRect"],
                "provenance": component["provenance"],
            })
    datasets = {
        "screen_contracts.csv": contracts,
        "screen_layout_index.csv": layout_rows,
        "screen_section_geometry.csv": section_rows,
        "screen_component_geometry.csv": component_rows,
        "database_tables.csv": snap["tables"],
        "database_columns.csv": snap["columns"],
        "database_constraints.csv": snap["constraints"],
        "database_indexes.csv": snap["indexes"],
    }
    for filename, rows in datasets.items():
        path = evidence / filename
        if not rows:
            path.write_text("", encoding="utf-8")
            continue
        with path.open("w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(f, fieldnames=list(rows[0].keys()))
            writer.writeheader()
            writer.writerows(rows)
    for filename, values in {
        "api_mappings.txt": snap["api_mappings"],
        "frontend_route_lines.txt": snap["frontend_route_lines"],
        "test_files.txt": snap["test_files"],
        "migration_files.txt": snap["migration_files"],
        "process_files.txt": snap["process_files"],
        "report_files.txt": snap["report_files"],
        "k8s_resources.txt": snap["k8s_resources"],
    }.items():
        (evidence / filename).write_text("\n".join(values) + "\n", encoding="utf-8")


def build_index(snap: dict, paths: list[Path]) -> Path:
    d = new_doc("통합 문서 목록 및 전수성 보고서", "PACK-FULL-INDEX-001", "19종 문서의 전수 포함 범위와 원천 수량을 확인합니다.", snap)
    add_overview(d, snap, [["설계 문서", len(paths), "요청 19종"], ["증적 파일", 17, "JSON/CSV/TXT"]])
    d.add_heading("2. 문서 구성", level=1)
    add_table(d, ["No", "문서", "파일", "전수 포함 대상"], [[i, p.stem.split("_", 1)[1], p.name, {
        1: "계약 액터 + 역할/권한 DB 행", 2: "process-step 유즈케이스", 3: "프로세스/단계/edge", 4: "sequence/handoff/readiness + 계약 로직", 5: "전문/계약 시나리오", 6: "계약 기반 요구사항", 7: "계약 규칙", 8: "전체 테이블 개념군", 9: "엔터티 + FK", 10: "제약 + 인덱스", 11: "전체 컬럼", 12: "전체 화면 계약", 13: "전체 route 구조", 14: "전체 UI 계약", 15: "전체 actor-process-step-route", 16: "계약 API + 소스 매핑", 17: "전체 테스트 파일", 18: "전체 route 테스트 시나리오", 19: "route별 기준 좌표 + Grid/Flex + 반응형 + 상태"}[i]] for i, p in enumerate(paths, 1)], [600, 1800, 3100, 3860], 7.5)
    d.add_heading("3. 전수성 판정", level=1)
    add_callout(d, "판정", "운영 원천에서 자동 추출 가능한 자산은 관련 문서에 전수 포함되었습니다. 다만 계약의 검증 플래그가 false인 항목과 원천에 서술되지 않은 사람 중심 업무 의미는 완료로 추정하지 않으며, 검토 필요 상태입니다.", "E9F6EC", GREEN)
    add_bullets(d, [
        "화면: 1,181개 계약과 1,000개 고유 route를 모두 포함했습니다.",
        "데이터: 344개 테이블, 4,243개 컬럼, 471개 인덱스, 2,870개 제약조건을 포함했습니다.",
        "소스: API 매핑 1,238줄, 테스트 파일 710개, 마이그레이션 254개를 증적에 보존했습니다.",
        "품질 점수 테이블은 현재 DB에 없어 점수 통과를 선언하지 않았습니다.",
    ])
    return save(d, "00_통합_문서_목록_및_전수성_보고서.docx")


def main() -> None:
    global SNAPSHOT, OUT, ZIP
    parser = argparse.ArgumentParser(description="Generate the full system design document pack.")
    parser.add_argument("--snapshot", type=Path, default=SNAPSHOT)
    parser.add_argument("--out", type=Path, default=OUT)
    parser.add_argument("--zip", dest="zip_path", type=Path)
    args = parser.parse_args()
    SNAPSHOT = args.snapshot.resolve()
    OUT = args.out.resolve()
    ZIP = (args.zip_path or OUT.parent / f"{OUT.name}.zip").resolve()
    snap = json.loads(SNAPSHOT.read_text(encoding="utf-8"))
    if OUT.exists():
        shutil.rmtree(OUT)
    OUT.mkdir(parents=True)
    contracts = decoded_contracts(snap)
    paths = build_docs(snap)
    index = build_index(snap, paths)
    export_evidence(snap, contracts)
    (OUT / "README.txt").write_text(
        f"{PROJECT}\nGenerated: {TODAY}\nSnapshot: {snap['captured_at']}\nDocuments: {len(paths) + 1}\nEvidence: evidence/\n",
        encoding="utf-8",
    )
    if ZIP.exists():
        ZIP.unlink()
    with zipfile.ZipFile(ZIP, "w", zipfile.ZIP_DEFLATED) as zf:
        for p in sorted(OUT.rglob("*")):
            if p.is_file():
                zf.write(p, p.relative_to(OUT.parent))
    print(json.dumps({"out": str(OUT), "zip": str(ZIP), "docx": len(list(OUT.glob('*.docx'))), "size": ZIP.stat().st_size, "index": str(index)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
